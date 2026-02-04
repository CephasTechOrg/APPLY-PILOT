"""
Resume extraction service - extracts text from PDF/DOCX and parses into canonical schema.
"""
import io
import json
import re
from typing import Any, Dict, Optional, Tuple

# PDF extraction - optional
try:
    import pdfplumber
    PDFPLUMBER_AVAILABLE = True
except ImportError:
    PDFPLUMBER_AVAILABLE = False

# DOCX extraction - optional
try:
    from docx import Document
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

from app.schemas.resume_content import (
    CanonicalResumeSchema,
    ContactInfo,
    EducationEntry,
    ExperienceEntry,
    ProfileSection,
    ProjectEntry,
    ResumeMeta,
    ResumeSections,
)


def extract_text_from_pdf(file_content: bytes) -> str:
    """
    Extract text from a PDF file.
    
    Args:
        file_content: Raw PDF file bytes
        
    Returns:
        Extracted text as string
    """
    if not PDFPLUMBER_AVAILABLE:
        raise RuntimeError("pdfplumber is not installed. Install with: pip install pdfplumber")
    
    text_parts = []
    
    with pdfplumber.open(io.BytesIO(file_content)) as pdf:
        for page in pdf.pages:
            page_text = page.extract_text()
            if page_text:
                text_parts.append(page_text)
    
    return "\n\n".join(text_parts)


def extract_text_from_docx(file_content: bytes) -> str:
    """
    Extract text from a DOCX file.
    
    Args:
        file_content: Raw DOCX file bytes
        
    Returns:
        Extracted text as string
    """
    if not DOCX_AVAILABLE:
        raise RuntimeError("python-docx is not installed. Install with: pip install python-docx")
    
    doc = Document(io.BytesIO(file_content))
    text_parts = []
    
    for paragraph in doc.paragraphs:
        if paragraph.text.strip():
            text_parts.append(paragraph.text)
    
    # Also extract text from tables
    for table in doc.tables:
        for row in table.rows:
            row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
            if row_text:
                text_parts.append(row_text)
    
    return "\n".join(text_parts)


def extract_text(file_content: bytes, content_type: str) -> str:
    """
    Extract text from a file based on content type.
    
    Args:
        file_content: Raw file bytes
        content_type: MIME type of the file
        
    Returns:
        Extracted text as string
        
    Raises:
        ValueError: If content type is not supported
    """
    if content_type == "application/pdf":
        return extract_text_from_pdf(file_content)
    elif content_type in [
        "application/msword",
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    ]:
        return extract_text_from_docx(file_content)
    else:
        raise ValueError(f"Unsupported content type: {content_type}")


# === Basic Parsing Helpers ===

def extract_email(text: str) -> Optional[str]:
    """Extract email address from text"""
    pattern = r'[a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,}'
    match = re.search(pattern, text)
    return match.group(0) if match else None


def extract_phone(text: str) -> Optional[str]:
    """Extract phone number from text"""
    patterns = [
        r'\+?1?[-.\s]?\(?[0-9]{3}\)?[-.\s]?[0-9]{3}[-.\s]?[0-9]{4}',
        r'\+?[0-9]{1,3}[-.\s]?[0-9]{3,4}[-.\s]?[0-9]{3,4}[-.\s]?[0-9]{3,4}',
    ]
    for pattern in patterns:
        match = re.search(pattern, text)
        if match:
            return match.group(0)
    return None


def extract_linkedin(text: str) -> Optional[str]:
    """Extract LinkedIn URL from text"""
    pattern = r'(?:https?://)?(?:www\.)?linkedin\.com/in/[\w-]+'
    match = re.search(pattern, text, re.IGNORECASE)
    return match.group(0) if match else None


def extract_github(text: str) -> Optional[str]:
    """Extract GitHub URL from text"""
    pattern = r'(?:https?://)?(?:www\.)?github\.com/[\w-]+'
    match = re.search(pattern, text, re.IGNORECASE)
    return match.group(0) if match else None


def basic_parse_resume(raw_text: str) -> Dict[str, Any]:
    """
    Basic rule-based parsing of resume text.
    This provides a starting point before AI enhancement.
    
    Args:
        raw_text: Extracted text from resume
        
    Returns:
        Dictionary with parsed resume data
    """
    lines = raw_text.split('\n')
    lines = [line.strip() for line in lines if line.strip()]
    
    # Try to extract name (usually first line or first bold/large text)
    full_name = lines[0] if lines else "Unknown"
    
    # Extract contact info
    email = extract_email(raw_text)
    phone = extract_phone(raw_text)
    linkedin = extract_linkedin(raw_text)
    portfolio = extract_github(raw_text)
    
    # Build basic schema
    result = {
        "meta": {
            "resumeId": None,
            "purpose": None,
            "industry": None,
            "language": "en",
            "tone": "professional"
        },
        "profile": {
            "fullName": full_name,
            "headline": None,
            "contact": {
                "email": email,
                "phone": phone,
                "location": None,
                "linkedin": linkedin,
                "portfolio": portfolio
            }
        },
        "sections": {
            "summary": None,
            "experience": [],
            "projects": [],
            "education": [],
            "skills": [],
            "certifications": [],
            "awards": []
        }
    }
    
    # Extract skills (look for common skill keywords)
    skill_keywords = [
        "Python", "JavaScript", "TypeScript", "React", "Node.js", "SQL",
        "Java", "C++", "C#", "Go", "Rust", "Ruby", "PHP", "Swift", "Kotlin",
        "AWS", "Azure", "GCP", "Docker", "Kubernetes", "Git", "Linux",
        "Machine Learning", "AI", "Data Science", "DevOps", "Agile"
    ]
    found_skills = []
    text_lower = raw_text.lower()
    for skill in skill_keywords:
        if skill.lower() in text_lower:
            found_skills.append(skill)
    result["sections"]["skills"] = found_skills
    
    return result


async def parse_resume_with_ai(
    raw_text: str,
    ai_service_func
) -> Dict[str, Any]:
    """
    Use AI to parse resume text into canonical schema.
    
    Args:
        raw_text: Extracted text from resume
        ai_service_func: Function to call AI service (can be sync or async)
        
    Returns:
        Dictionary with parsed resume data in canonical schema format
    """
    prompt = f"""Parse the following resume text into a structured JSON format.

Return ONLY valid JSON with this exact structure:
{{
  "meta": {{
    "resumeId": null,
    "purpose": null,
    "industry": "<detected industry or null>",
    "language": "en",
    "tone": "professional"
  }},
  "profile": {{
    "fullName": "<full name>",
    "headline": "<professional headline or null>",
    "contact": {{
      "email": "<email or null>",
      "phone": "<phone or null>",
      "location": "<location or null>",
      "linkedin": "<linkedin url or null>",
      "portfolio": "<portfolio/github url or null>"
    }}
  }},
  "sections": {{
    "summary": "<professional summary or null>",
    "experience": [
      {{
        "company": "<company name>",
        "role": "<job title>",
        "location": "<location or null>",
        "startDate": "<start date>",
        "endDate": "<end date or Present>",
        "bullets": ["<achievement 1>", "<achievement 2>"]
      }}
    ],
    "projects": [
      {{
        "name": "<project name>",
        "description": "<description>",
        "technologies": ["<tech1>", "<tech2>"]
      }}
    ],
    "education": [
      {{
        "institution": "<school name>",
        "degree": "<degree>",
        "field": "<field of study>",
        "startDate": "<start>",
        "endDate": "<end>"
      }}
    ],
    "skills": ["<skill1>", "<skill2>"],
    "certifications": [
      {{
        "name": "<cert name>",
        "issuer": "<issuer>",
        "date": "<date or null>"
      }}
    ],
    "awards": []
  }}
}}

Resume text:
---
{raw_text[:8000]}
---

Return ONLY the JSON, no markdown, no explanation."""

    try:
        # Call AI service (supports both sync and async)
        import asyncio
        import inspect
        
        if inspect.iscoroutinefunction(ai_service_func):
            response = await ai_service_func(prompt)
        else:
            # Sync function - call directly
            response, _ = ai_service_func(prompt)
        
        # Handle tuple response (response, tokens)
        if isinstance(response, tuple):
            response = response[0]
        
        # Try to extract JSON from response
        json_text = response.strip()
        if json_text.startswith("```"):
            # Remove markdown code blocks
            json_text = re.sub(r'^```(?:json)?\n?', '', json_text)
            json_text = re.sub(r'\n?```$', '', json_text)
        
        parsed = json.loads(json_text)
        return parsed
        
    except Exception as e:
        # Fall back to basic parsing
        print(f"AI parsing failed: {e}, falling back to basic parsing")
        return basic_parse_resume(raw_text)


def validate_resume_schema(data: Dict[str, Any]) -> Tuple[bool, Optional[str]]:
    """
    Validate that parsed data conforms to canonical schema.
    
    Args:
        data: Parsed resume data
        
    Returns:
        Tuple of (is_valid, error_message)
    """
    try:
        # Attempt to construct the Pydantic model
        profile_data = data.get("profile", {})
        if not profile_data.get("fullName"):
            return False, "Missing required field: profile.fullName"
        
        CanonicalResumeSchema(
            meta=ResumeMeta(**data.get("meta", {})),
            profile=ProfileSection(
                fullName=profile_data.get("fullName"),
                headline=profile_data.get("headline"),
                contact=ContactInfo(**profile_data.get("contact", {}))
            ),
            sections=ResumeSections(**data.get("sections", {}))
        )
        return True, None
        
    except Exception as e:
        return False, str(e)
