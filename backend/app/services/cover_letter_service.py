"""Cover Letter rendering and export service."""

import io
import json
from datetime import datetime
from pathlib import Path
from typing import Optional

from jinja2 import Template

# PDF export - WeasyPrint requires GTK libraries on Windows
try:
    from weasyprint import HTML
    WEASYPRINT_AVAILABLE = True
except (ImportError, OSError):
    HTML = None
    WEASYPRINT_AVAILABLE = False

# DOCX export
try:
    from docx import Document
    from docx.shared import Pt, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    DOCX_AVAILABLE = True
except ImportError:
    Document = None
    DOCX_AVAILABLE = False


TEMPLATES_DIR = Path(__file__).parent.parent.parent / "templates" / "cover_letter"


def get_available_templates() -> list[dict]:
    """Get list of available cover letter templates."""
    templates = []
    
    if not TEMPLATES_DIR.exists():
        return templates
    
    for template_dir in TEMPLATES_DIR.iterdir():
        if template_dir.is_dir():
            config_path = template_dir / "config.json"
            if config_path.exists():
                with open(config_path, "r", encoding="utf-8") as f:
                    config = json.load(f)
                    templates.append({
                        "slug": template_dir.name,
                        "name": config.get("name", template_dir.name.title()),
                        "description": config.get("description", ""),
                        "supported_purposes": config.get("supported_purposes", []),
                        "default_tokens": config.get("default_tokens", {}),
                        "features": config.get("features", []),
                    })
    
    return templates


def get_template(slug: str) -> Optional[dict]:
    """Get a specific template by slug."""
    template_dir = TEMPLATES_DIR / slug
    if not template_dir.exists():
        return None
    
    config_path = template_dir / "config.json"
    html_path = template_dir / "index.html"
    css_path = template_dir / "styles.css"
    
    if not all(p.exists() for p in [config_path, html_path, css_path]):
        return None
    
    with open(config_path, "r", encoding="utf-8") as f:
        config = json.load(f)
    
    with open(html_path, "r", encoding="utf-8") as f:
        html_content = f.read()
    
    with open(css_path, "r", encoding="utf-8") as f:
        css_content = f.read()
    
    return {
        "slug": slug,
        "name": config.get("name", slug.title()),
        "description": config.get("description", ""),
        "html_content": html_content,
        "css_content": css_content,
        "config": config,
    }


def resolve_design_tokens(tokens: Optional[dict]) -> dict:
    """Resolve design tokens to CSS classes."""
    default_tokens = {
        "fontFamily": "Inter",
        "spacing": "comfortable",
        "accentColor": "neutral",
    }
    
    if tokens:
        default_tokens.update(tokens)
    
    # Map to CSS class names
    font_map = {
        "Inter": "font-inter",
        "Roboto": "font-roboto",
        "Georgia": "font-georgia",
    }
    
    spacing_map = {
        "compact": "spacing-compact",
        "comfortable": "spacing-comfortable",
    }
    
    accent_map = {
        "neutral": "accent-neutral",
        "blue": "accent-blue",
        "green": "accent-green",
    }
    
    classes = []
    classes.append(font_map.get(default_tokens["fontFamily"], "font-inter"))
    classes.append(spacing_map.get(default_tokens["spacing"], "spacing-comfortable"))
    classes.append(accent_map.get(default_tokens["accentColor"], "accent-neutral"))
    
    return {
        "classes": " ".join(classes),
        "tokens": default_tokens,
    }


def render_cover_letter_html(
    content: dict,
    template_slug: str = "formal",
    design_tokens: Optional[dict] = None,
    profile: Optional[dict] = None,
) -> str:
    """Render a cover letter to HTML using the specified template."""
    template_data = get_template(template_slug)
    if not template_data:
        # Fallback to formal template
        template_data = get_template("formal")
        if not template_data:
            raise ValueError(f"Template '{template_slug}' not found and no fallback available")
    
    # Resolve design tokens
    token_result = resolve_design_tokens(design_tokens)
    
    # Add token classes to body tag in CSS
    css_content = template_data["css_content"]
    
    # Prepare template context
    meta = content.get("meta", {})
    letter_content = content.get("content", {})
    
    # Format current date
    current_date = datetime.now().strftime("%B %d, %Y")
    
    context = {
        "css": css_content,
        "meta": meta,
        "content": letter_content,
        "profile": profile or {},
        "current_date": current_date,
        "token_classes": token_result["classes"],
    }
    
    # Render HTML template
    html_template = Template(template_data["html_content"])
    rendered_html = html_template.render(**context)
    
    # Inject token classes into body tag
    rendered_html = rendered_html.replace(
        "<body>",
        f'<body class="{token_result["classes"]}">'
    )
    
    return rendered_html


def export_to_pdf(
    content: dict,
    template_slug: str = "formal",
    design_tokens: Optional[dict] = None,
    profile: Optional[dict] = None,
) -> bytes:
    """Export cover letter to PDF."""
    if not WEASYPRINT_AVAILABLE:
        raise RuntimeError("WeasyPrint is not installed. Install with: pip install weasyprint")
    
    html_content = render_cover_letter_html(content, template_slug, design_tokens, profile)
    
    # Generate PDF
    html_doc = HTML(string=html_content)
    pdf_bytes = html_doc.write_pdf()
    
    return pdf_bytes


def export_to_docx(
    content: dict,
    profile: Optional[dict] = None,
) -> bytes:
    """Export cover letter to DOCX."""
    if not DOCX_AVAILABLE:
        raise RuntimeError("python-docx is not installed. Install with: pip install python-docx")
    
    doc = Document()
    
    meta = content.get("meta", {})
    letter_content = content.get("content", {})
    
    # Add sender name
    sender_name = letter_content.get("sender_name") or (profile or {}).get("fullName", "")
    if sender_name:
        heading = doc.add_heading(sender_name, level=1)
        heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
    
    # Add contact info
    contact = (profile or {}).get("contact", {})
    if contact:
        contact_parts = []
        if contact.get("email"):
            contact_parts.append(contact["email"])
        if contact.get("phone"):
            contact_parts.append(contact["phone"])
        if contact.get("location"):
            contact_parts.append(contact["location"])
        if contact_parts:
            contact_para = doc.add_paragraph(" | ".join(contact_parts))
            contact_para.style = "Subtitle"
    
    # Add date
    date_para = doc.add_paragraph(datetime.now().strftime("%B %d, %Y"))
    date_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    
    doc.add_paragraph()  # Spacer
    
    # Add recipient info
    if meta.get("company_name"):
        doc.add_paragraph(meta["company_name"])
    if meta.get("job_title"):
        doc.add_paragraph(f"Re: {meta['job_title']}")
    
    doc.add_paragraph()  # Spacer
    
    # Salutation
    salutation = letter_content.get("salutation", "Dear Hiring Manager,")
    doc.add_paragraph(salutation)
    
    doc.add_paragraph()  # Spacer
    
    # Opening
    if letter_content.get("opening"):
        doc.add_paragraph(letter_content["opening"])
    
    # Body paragraphs
    for paragraph in letter_content.get("body", []):
        doc.add_paragraph(paragraph)
    
    # Closing
    if letter_content.get("closing"):
        doc.add_paragraph(letter_content["closing"])
    
    doc.add_paragraph()  # Spacer
    
    # Signature
    signature = letter_content.get("signature", "Sincerely,")
    doc.add_paragraph(signature)
    
    doc.add_paragraph()  # Space for signature
    doc.add_paragraph()
    
    sender_sig = letter_content.get("sender_name") or (profile or {}).get("fullName", "")
    if sender_sig:
        sig_para = doc.add_paragraph(sender_sig)
        sig_para.runs[0].bold = True
    
    # Save to bytes
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    
    return buffer.getvalue()
