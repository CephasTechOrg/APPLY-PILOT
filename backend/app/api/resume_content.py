"""
API endpoints for resume content extraction and management.
"""
import asyncio
from typing import Optional

from fastapi import APIRouter, BackgroundTasks, Depends, HTTPException, status
from sqlalchemy.orm import Session

from app.api.deps import get_current_user
from app.core.database import get_db, SessionLocal
from app.core.storage import download_resume_file
from app.models.resume import Resume
from app.models.resume_content import ExtractionStatus, ResumeContent
from app.models.user import User
from app.schemas.resume_content import (
    ExtractResumeRequest,
    ExtractResumeResponse,
    ResumeContentCreate,
    ResumeContentResponse,
    ResumeContentUpdate,
)
from app.schemas.template import ExportResumeRequest, RenderResumeRequest, RenderResumeResponse
from app.services.extraction_service import (
    basic_parse_resume,
    extract_text,
    parse_resume_with_ai,
    validate_resume_schema,
)
from app.services.template_service import render_resume_html, resolve_design_tokens

router = APIRouter()


def _run_extraction_sync(
    resume_id: int,
    use_ai: bool,
):
    """
    Background task to extract and parse resume content.
    Uses its own database session to avoid session-in-different-thread issues.
    """
    from app.services.ai_service import _call_deepseek
    
    # Create a fresh database session for this background task
    db = SessionLocal()
    
    try:
        # Get resume
        resume = db.query(Resume).filter(Resume.id == resume_id).first()
        if not resume:
            return
        
        # Get or create content record
        content = db.query(ResumeContent).filter(ResumeContent.resume_id == resume_id).first()
        if not content:
            content = ResumeContent(resume_id=resume_id)
            db.add(content)
        
        content.extraction_status = ExtractionStatus.PROCESSING.value
        db.commit()
        
        try:
            # Download file from storage
            file_bytes = download_resume_file(resume.storage_path)
            
            # Extract text
            raw_text = extract_text(file_bytes, resume.content_type)
            
            if not raw_text.strip():
                raise ValueError("No text could be extracted from the resume")
            
            # Parse to schema
            if use_ai:
                # Run async function in sync context
                loop = asyncio.new_event_loop()
                asyncio.set_event_loop(loop)
                try:
                    parsed_data = loop.run_until_complete(parse_resume_with_ai(raw_text, _call_deepseek))
                finally:
                    loop.close()
            else:
                parsed_data = basic_parse_resume(raw_text)
            
            # Validate schema
            is_valid, error = validate_resume_schema(parsed_data)
            if not is_valid:
                content.extraction_status = ExtractionStatus.FAILED.value
                content.extraction_error = f"Schema validation failed: {error}"
                db.commit()
                return
            
            # Save parsed data
            content.structured_data = parsed_data
            content.extraction_status = ExtractionStatus.COMPLETED.value
            content.extraction_error = None
            
            # Extract metadata
            meta = parsed_data.get("meta", {})
            content.purpose = meta.get("purpose")
            content.industry = meta.get("industry")
            content.language = meta.get("language", "en")
            content.tone = meta.get("tone", "professional")
            
            db.commit()
            
        except Exception as e:
            content.extraction_status = ExtractionStatus.FAILED.value
            content.extraction_error = str(e)
            db.commit()
    finally:
        db.close()


@router.post("/{resume_id}/extract", response_model=ExtractResumeResponse)
async def extract_resume_content(
    resume_id: int,
    background_tasks: BackgroundTasks,
    use_ai: bool = True,
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db),
):
    """
    Trigger extraction of resume content from uploaded file.
    Runs in background and updates ResumeContent record.
    """
    # Verify resume belongs to user
    resume = (
        db.query(Resume)
        .filter(Resume.id == resume_id, Resume.user_id == current_user.id)
        .first()
    )
    if not resume:
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND,
            detail="Resume not found"
        )
    
    # Check if extraction is already in progress
    existing = db.query(ResumeContent).filter(ResumeContent.resume_id == resume_id).first()
    if existing and existing.extraction_status == ExtractionStatus.PROCESSING.value:
        return ExtractResumeResponse(
            resume_id=resume_id,
            status="processing",
            message="Extraction already in progress"
        )
    
    # Create or reset content record
    if not existing:
        existing = ResumeContent(
            resume_id=resume_id,
            extraction_status=ExtractionStatus.PENDING.value
        )
        db.add(existing)
        db.commit()
    
    # Queue background extraction
    background_tasks.add_task(_run_extraction_sync, resume_id, use_ai)
    
    return ExtractResumeResponse(
        resume_id=resume_id,
        status="pending",
        message="Extraction started. Check back for results."
    )


@router.get("/{resume_id}/content", response_model=ResumeContentResponse)
async def get_resume_content(
    resume_id: int,
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db),
):
    """
    Get parsed resume content for a resume.
    """
    # Verify resume belongs to user
    resume = (
        db.query(Resume)
        .filter(Resume.id == resume_id, Resume.user_id == current_user.id)
        .first()
    )
    if not resume:
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND,
            detail="Resume not found"
        )
    
    content = db.query(ResumeContent).filter(ResumeContent.resume_id == resume_id).first()
    if not content:
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND,
            detail="Resume content not extracted yet. Call /extract first."
        )
    
    return content


@router.put("/{resume_id}/content", response_model=ResumeContentResponse)
async def update_resume_content(
    resume_id: int,
    payload: ResumeContentUpdate,
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db),
):
    """
    Update parsed resume content (manual edits).
    """
    resume = (
        db.query(Resume)
        .filter(Resume.id == resume_id, Resume.user_id == current_user.id)
        .first()
    )
    if not resume:
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND,
            detail="Resume not found"
        )
    
    content = db.query(ResumeContent).filter(ResumeContent.resume_id == resume_id).first()
    if not content:
        # Create new content record
        content = ResumeContent(resume_id=resume_id)
        db.add(content)
    
    # Update fields
    if payload.structured_data is not None:
        # Validate new data
        is_valid, error = validate_resume_schema(payload.structured_data)
        if not is_valid:
            raise HTTPException(
                status_code=status.HTTP_400_BAD_REQUEST,
                detail=f"Invalid resume schema: {error}"
            )
        content.structured_data = payload.structured_data
        content.extraction_status = ExtractionStatus.COMPLETED.value
    
    if payload.purpose is not None:
        content.purpose = payload.purpose
    if payload.industry is not None:
        content.industry = payload.industry
    if payload.language is not None:
        content.language = payload.language
    if payload.tone is not None:
        content.tone = payload.tone
    
    db.commit()
    db.refresh(content)
    
    return content


@router.post("/{resume_id}/preview", response_model=RenderResumeResponse)
async def preview_resume(
    resume_id: int,
    payload: RenderResumeRequest,
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db),
):
    """
    Render resume with selected template for preview.
    """
    # Get resume content
    resume = (
        db.query(Resume)
        .filter(Resume.id == resume_id, Resume.user_id == current_user.id)
        .first()
    )
    if not resume:
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND,
            detail="Resume not found"
        )
    
    content = db.query(ResumeContent).filter(ResumeContent.resume_id == resume_id).first()
    if not content or not content.structured_data:
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail="Resume content not available. Extract or create content first."
        )
    
    # Resolve tokens
    tokens = resolve_design_tokens(payload.design_tokens, payload.template_slug)
    
    # Render
    try:
        html = render_resume_html(
            template_slug=payload.template_slug,
            resume_data=content.structured_data,
            design_tokens=payload.design_tokens,
            purpose=payload.purpose or content.purpose
        )
    except ValueError as e:
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail=str(e)
        )
    
    return RenderResumeResponse(
        html=html,
        template_slug=payload.template_slug,
        tokens_used=tokens
    )


@router.post("/{resume_id}/export")
async def export_resume(
    resume_id: int,
    payload: ExportResumeRequest,
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db),
):
    """
    Export resume as PDF or DOCX.
    """
    from fastapi.responses import Response
    from app.services.template_service import render_resume_for_export
    
    # Get resume content
    resume = (
        db.query(Resume)
        .filter(Resume.id == resume_id, Resume.user_id == current_user.id)
        .first()
    )
    if not resume:
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND,
            detail="Resume not found"
        )
    
    content = db.query(ResumeContent).filter(ResumeContent.resume_id == resume_id).first()
    if not content or not content.structured_data:
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail="Resume content not available."
        )
    
    # Render HTML
    html = render_resume_for_export(
        template_slug=payload.template_slug,
        resume_data=content.structured_data,
        design_tokens=payload.design_tokens,
        purpose=content.purpose,
        page_size=payload.page_size
    )
    
    if payload.format == "pdf":
        # Generate PDF using WeasyPrint
        try:
            from weasyprint import HTML
            pdf_bytes = HTML(string=html).write_pdf()
            
            filename = f"{resume.title.replace(' ', '_')}.pdf"
            return Response(
                content=pdf_bytes,
                media_type="application/pdf",
                headers={
                    "Content-Disposition": f'attachment; filename="{filename}"'
                }
            )
        except Exception as e:
            raise HTTPException(
                status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
                detail=f"PDF generation failed: {e}"
            )
    
    elif payload.format == "docx":
        # Generate DOCX using python-docx
        try:
            from docx import Document
            from docx.shared import Inches, Pt
            from io import BytesIO
            import re
            
            doc = Document()
            
            # Extract text from HTML (simple approach)
            # In production, use a proper HTML parser
            profile = content.structured_data.get("profile", {})
            sections = content.structured_data.get("sections", {})
            
            # Add name
            name_para = doc.add_heading(profile.get("fullName", ""), 0)
            
            # Add contact info
            contact = profile.get("contact", {})
            contact_parts = []
            if contact.get("email"):
                contact_parts.append(contact["email"])
            if contact.get("phone"):
                contact_parts.append(contact["phone"])
            if contact.get("location"):
                contact_parts.append(contact["location"])
            if contact_parts:
                doc.add_paragraph(" | ".join(contact_parts))
            
            # Add sections
            if sections.get("summary"):
                doc.add_heading("Summary", level=1)
                doc.add_paragraph(sections["summary"])
            
            if sections.get("experience"):
                doc.add_heading("Experience", level=1)
                for exp in sections["experience"]:
                    doc.add_heading(f"{exp.get('role', '')} at {exp.get('company', '')}", level=2)
                    doc.add_paragraph(f"{exp.get('startDate', '')} - {exp.get('endDate', 'Present')}")
                    for bullet in exp.get("bullets", []):
                        doc.add_paragraph(bullet, style="List Bullet")
            
            if sections.get("education"):
                doc.add_heading("Education", level=1)
                for edu in sections["education"]:
                    doc.add_paragraph(f"{edu.get('institution', '')} - {edu.get('degree', '')} in {edu.get('field', '')}")
            
            if sections.get("skills"):
                doc.add_heading("Skills", level=1)
                doc.add_paragraph(", ".join(sections["skills"]))
            
            # Save to bytes
            buffer = BytesIO()
            doc.save(buffer)
            buffer.seek(0)
            
            filename = f"{resume.title.replace(' ', '_')}.docx"
            return Response(
                content=buffer.getvalue(),
                media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                headers={
                    "Content-Disposition": f'attachment; filename="{filename}"'
                }
            )
        except Exception as e:
            raise HTTPException(
                status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
                detail=f"DOCX generation failed: {e}"
            )
    
    else:
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail=f"Unsupported format: {payload.format}"
        )
